import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_CN_SIZES = {
    "初号": 84, "小初": 72, "一号": 52, "小一": 48,
    "二号": 44, "小二": 36, "三号": 32, "小三": 30,
    "四号": 28, "小四": 24, "五号": 21, "小五": 18,
}


class DocFormatter:
    def __init__(self, docx_path):
        self.doc = Document(docx_path)

    # ══════════════════════════════════════════════════════════════════
    #  extract — 供 LLM 阅读理解
    # ══════════════════════════════════════════════════════════════════

    def extract_paragraphs(self, max_chars=80):
        result = []
        for i, p in enumerate(self.doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue
            result.append({
                "idx": i,
                "style": p.style.name,
                "text": text[:max_chars] + ("..." if len(text) > max_chars else ""),
            })
        return result

    def extract_tables(self, max_cell_chars=30):
        result = []
        for i, table in enumerate(self.doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            # 采样第一行作为表头
            header = []
            if rows > 0:
                header = [c.text.strip()[:max_cell_chars]
                          for c in table.rows[0].cells]
            # 当前行高（取第一个单元格的行高）
            row_h = None
            if rows > 0:
                tr = table.rows[0]._tr
                trPr = tr.find(f"{{{NS}}}trPr")
                if trPr is not None:
                    trHeight = trPr.find(f"{{{NS}}}trHeight")
                    if trHeight is not None:
                        val = trHeight.get(f"{{{NS}}}val")
                        if val:
                            row_h = round(int(val) / 360000, 2)  # EMU → cm
            # 列宽
            col_widths = []
            if cols > 0:
                for col in table.columns:
                    w = col.width
                    if w is not None and w > 0:
                        col_widths.append(round(w / 360000, 2))
            result.append({
                "idx": i,
                "rows": rows,
                "cols": cols,
                "header": header,
                "row_height_cm": row_h,
                "col_widths_cm": col_widths,
            })
        return result

    def extract_headers_footers(self):
        result = []
        for i, section in enumerate(self.doc.sections):
            info = {"section_idx": i}
            # 页眉
            if section.header and not section.header.is_linked_to_previous:
                texts = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
                info["header_text"] = " | ".join(texts) if texts else ""
                if section.header.paragraphs:
                    p0 = section.header.paragraphs[0]
                    if p0.runs:
                        r = p0.runs[0]
                        info["header_font"] = r.font.name or ""
                        info["header_size_pt"] = round(r.font.size / 12700, 1) if r.font.size else None
                    info["header_align"] = str(p0.alignment) if p0.alignment else ""
            # 页脚
            if section.footer and not section.footer.is_linked_to_previous:
                texts = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
                info["footer_text"] = " | ".join(texts) if texts else ""
                # 检查是否有页码域
                for p in section.footer.paragraphs:
                    for r in p.runs:
                        if r._element.findall(f".//{{{NS}}}fldChar"):
                            info["footer_has_page_num"] = True
                            break
            result.append(info)
        return result

    # ══════════════════════════════════════════════════════════════════
    #  execute — 应用 LLM 返回的格式
    # ══════════════════════════════════════════════════════════════════

    def execute(self, result):
        """
        result 是列表，每个元素可以是：
        - 段落格式: {"idx": int, "font": ..., ...}
        - 表格格式: {"table_idx": int, ...}
        - 页眉页脚格式: {"section_idx": int, ...}
        """
        for item in result:
            if "table_idx" in item:
                self._apply_table(item)
            elif "section_idx" in item:
                self._apply_header_footer(item)
            elif "idx" in item:
                idx = item["idx"]
                if idx < len(self.doc.paragraphs):
                    self._apply(self.doc.paragraphs[idx], item)

    def save_as(self, output_path):
        self.doc.save(output_path)

    # ══════════════════════════════════════════════════════════════════
    #  表格操作
    # ══════════════════════════════════════════════════════════════════

    def _apply_table(self, fmt):
        idx = fmt.get("table_idx", 0)
        if idx >= len(self.doc.tables):
            return
        table = self.doc.tables[idx]

        # 行高
        if "row_height_cm" in fmt:
            h = Cm(fmt["row_height_cm"])
            for row in table.rows:
                tr = row._tr
                trPr = tr.find(f"{{{NS}}}trPr")
                if trPr is None:
                    trPr = etree.SubElement(tr, f"{{{NS}}}trPr")
                trHeight = trPr.find(f"{{{NS}}}trHeight")
                if trHeight is None:
                    trHeight = etree.SubElement(trPr, f"{{{NS}}}trHeight")
                trHeight.set(f"{{{NS}}}val", str(int(h)))
                trHeight.set(f"{{{NS}}}hRule", "exact")

        # 列宽
        if "col_widths_cm" in fmt:
            widths = fmt["col_widths_cm"]
            for i, col in enumerate(table.columns):
                if i < len(widths):
                    col.width = Cm(widths[i])

        # 单元格字体/字号/对齐
        font_name = fmt.get("font")
        size_pt = fmt.get("size_pt")
        size_hint = fmt.get("size_hint")
        if size_hint and size_hint in _CN_SIZES:
            size_pt = _CN_SIZES[size_hint] / 2
        align = fmt.get("align")
        header_bold = fmt.get("header_bold")

        for row_i, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 对齐
                    if align and align in ALIGN_MAP:
                        para.alignment = ALIGN_MAP[align]
                    # 字体
                    for run in para.runs:
                        rpr = run._element.find(f"{{{NS}}}rPr")
                        if rpr is None:
                            rpr = etree.SubElement(run._element, f"{{{NS}}}rPr")
                        if font_name:
                            rfonts = rpr.find(f"{{{NS}}}rFonts")
                            if rfonts is None:
                                rfonts = etree.SubElement(rpr, f"{{{NS}}}rFonts")
                            rfonts.set(f"{{{NS}}}ascii", font_name)
                            rfonts.set(f"{{{NS}}}hAnsi", font_name)
                            rfonts.set(f"{{{NS}}}eastAsia", font_name)
                        if size_pt:
                            sz_val = str(int(size_pt * 2))
                            sz = rpr.find(f"{{{NS}}}sz")
                            if sz is None:
                                sz = etree.SubElement(rpr, f"{{{NS}}}sz")
                            sz.set(f"{{{NS}}}val", sz_val)
                        # 表头加粗
                        if header_bold and row_i == 0:
                            b = rpr.find(f"{{{NS}}}b")
                            if b is None:
                                b = etree.SubElement(rpr, f"{{{NS}}}b")
                            b.attrib.pop(f"{{{NS}}}val", None)

    # ══════════════════════════════════════════════════════════════════
    #  页眉页脚操作
    # ══════════════════════════════════════════════════════════════════

    def _apply_header_footer(self, fmt):
        idx = fmt.get("section_idx", 0)
        if idx >= len(self.doc.sections):
            return
        section = self.doc.sections[idx]

        # 页眉
        if any(k.startswith("header_") for k in fmt):
            header = section.header
            header.is_linked_to_previous = False
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            if "header_text" in fmt:
                para.clear()
                para.add_run(fmt["header_text"])
            self._fmt_runs(para, fmt, "header_")
            if "header_align" in fmt and fmt["header_align"] in ALIGN_MAP:
                para.alignment = ALIGN_MAP[fmt["header_align"]]

        # 页脚
        if any(k.startswith("footer_") for k in fmt):
            footer = section.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if "footer_text" in fmt:
                para.clear()
                para.add_run(fmt["footer_text"])
            self._fmt_runs(para, fmt, "footer_")
            if "footer_align" in fmt and fmt["footer_align"] in ALIGN_MAP:
                para.alignment = ALIGN_MAP[fmt["footer_align"]]
            # 页码
            if fmt.get("footer_page_num"):
                if fmt.get("footer_text"):
                    para.add_run(" - ")
                self._insert_page_num(para)

    def _fmt_runs(self, para, fmt, prefix):
        font_name = fmt.get(f"{prefix}font")
        size_pt = fmt.get(f"{prefix}size_pt")
        size_hint = fmt.get(f"{prefix}size_hint")
        if size_hint and size_hint in _CN_SIZES:
            size_pt = _CN_SIZES[size_hint] / 2
        for run in para.runs:
            rpr = run._element.find(f"{{{NS}}}rPr")
            if rpr is None:
                rpr = etree.SubElement(run._element, f"{{{NS}}}rPr")
            if font_name:
                rfonts = rpr.find(f"{{{NS}}}rFonts")
                if rfonts is None:
                    rfonts = etree.SubElement(rpr, f"{{{NS}}}rFonts")
                rfonts.set(f"{{{NS}}}ascii", font_name)
                rfonts.set(f"{{{NS}}}hAnsi", font_name)
                rfonts.set(f"{{{NS}}}eastAsia", font_name)
            if size_pt:
                sz_val = str(int(size_pt * 2))
                sz = rpr.find(f"{{{NS}}}sz")
                if sz is None:
                    sz = etree.SubElement(rpr, f"{{{NS}}}sz")
                sz.set(f"{{{NS}}}val", sz_val)

    def _insert_page_num(self, para):
        """插入页码域"""
        run = para.add_run()
        fld = etree.SubElement(run._element, f"{{{NS}}}fldChar")
        fld.set(f"{{{NS}}}fldCharType", "begin")
        run2 = para.add_run()
        instr = etree.SubElement(run2._element, f"{{{NS}}}instrText")
        instr.set(f"{{{NS}}}space", "preserve")
        instr.text = " PAGE "
        run3 = para.add_run()
        fld2 = etree.SubElement(run3._element, f"{{{NS}}}fldChar")
        fld2.set(f"{{{NS}}}fldCharType", "end")

    # ══════════════════════════════════════════════════════════════════
    #  段落操作（原有逻辑）
    # ══════════════════════════════════════════════════════════════════

    def _apply(self, para, fmt):
        font_name = fmt.get("font")
        size_pt = fmt.get("size_pt")
        size_hint = fmt.get("size_hint")
        if size_hint and size_hint in _CN_SIZES:
            size_pt = _CN_SIZES[size_hint] / 2

        for run in para.runs:
            rpr = run._element.find(f"{{{NS}}}rPr")
            if rpr is None:
                rpr = etree.SubElement(run._element, f"{{{NS}}}rPr")

            if font_name or size_pt:
                rfonts = rpr.find(f"{{{NS}}}rFonts")
                if rfonts is None:
                    rfonts = etree.SubElement(rpr, f"{{{NS}}}rFonts")
                if font_name:
                    rfonts.set(f"{{{NS}}}ascii", font_name)
                    rfonts.set(f"{{{NS}}}hAnsi", font_name)
                    rfonts.set(f"{{{NS}}}eastAsia", font_name)
                    rfonts.set(f"{{{NS}}}cs", font_name)
                if size_pt:
                    sz_val = str(int(size_pt * 2))
                    sz = rpr.find(f"{{{NS}}}sz")
                    if sz is None:
                        sz = etree.SubElement(rpr, f"{{{NS}}}sz")
                    sz.set(f"{{{NS}}}val", sz_val)
                    szCs = rpr.find(f"{{{NS}}}szCs")
                    if szCs is None:
                        szCs = etree.SubElement(rpr, f"{{{NS}}}szCs")
                    szCs.set(f"{{{NS}}}val", sz_val)

            if "bold" in fmt:
                b = rpr.find(f"{{{NS}}}b")
                if b is None:
                    b = etree.SubElement(rpr, f"{{{NS}}}b")
                if fmt["bold"]:
                    b.attrib.pop(f"{{{NS}}}val", None)
                else:
                    b.set(f"{{{NS}}}val", "0")

            if "italic" in fmt:
                i_elem = rpr.find(f"{{{NS}}}i")
                if i_elem is None:
                    i_elem = etree.SubElement(rpr, f"{{{NS}}}i")
                if fmt["italic"]:
                    i_elem.attrib.pop(f"{{{NS}}}val", None)
                else:
                    i_elem.set(f"{{{NS}}}val", "0")

            if "color" in fmt:
                hex_str = fmt["color"].lstrip("#")
                color_elem = rpr.find(f"{{{NS}}}color")
                if color_elem is None:
                    color_elem = etree.SubElement(rpr, f"{{{NS}}}color")
                color_elem.set(f"{{{NS}}}val", hex_str)

        if "align" in fmt and fmt["align"] in ALIGN_MAP:
            para.alignment = ALIGN_MAP[fmt["align"]]

        pf = para.paragraph_format
        if "indent_chars" in fmt:
            pPr = para._element.find(f"{{{NS}}}pPr")
            if pPr is None:
                pPr = etree.SubElement(para._element, f"{{{NS}}}pPr")
            ind = pPr.find(f"{{{NS}}}ind")
            if ind is None:
                ind = etree.SubElement(pPr, f"{{{NS}}}ind")
            ind.set(f"{{{NS}}}firstLineChars", str(int(float(fmt["indent_chars"]) * 100)))
            ind.attrib.pop(f"{{{NS}}}firstLine", None)
        if "line_spacing" in fmt:
            pf.line_spacing = float(fmt["line_spacing"])
        if "line_spacing_pt" in fmt:
            pf.line_spacing = Pt(float(fmt["line_spacing_pt"]))
        if "space_before_pt" in fmt:
            pf.space_before = Pt(float(fmt["space_before_pt"]))
        if "space_after_pt" in fmt:
            pf.space_after = Pt(float(fmt["space_after_pt"]))
